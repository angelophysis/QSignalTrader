"""
Gera documentação completa do QSignalTrader em formato DOCX.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime


def _set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table


def _code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F0F0')
    p._element.get_or_add_pPr().append(shd)
    return p


def _h1(doc, text):
    return doc.add_heading(text, level=1)


def _h2(doc, text):
    return doc.add_heading(text, level=2)


def _h3(doc, text):
    return doc.add_heading(text, level=3)


def _p(doc, text):
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)
    return p


def _bullet(doc, text):
    return doc.add_paragraph(text, style='List Bullet')


def _note(doc, text):
    p = doc.add_paragraph(text, style='Intense Quote')
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    return p


def criar_documentacao():
    doc = Document()

    # ── Estilos globais ──
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(4)

    # ── Capa ──
    for _ in range(6):
        doc.add_paragraph()
    capa_titulo = doc.add_paragraph()
    capa_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = capa_titulo.add_run('QSignalTrader')
    run.bold = True
    run.font.size = Pt(32)
    run.font.color.rgb = RGBColor(0x0D, 0x11, 0x17)

    capa_sub = doc.add_paragraph()
    capa_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = capa_sub.add_run(
        'Motor Quantitativo Multi-Ativo para\n'
        'Análise Direcional, Volatilidade e Decisão Operacional'
    )
    run_sub.font.size = Pt(14)
    run_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    linha = doc.add_paragraph()
    linha.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_linha = linha.add_run('─' * 60)
    run_linha.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

    doc.add_paragraph()
    info = doc.add_paragraph()
    info.alignment = WD_ALIGN_PARAGRAPH.CENTER
    info.add_run(f'Versão: 1.0.0\nGerado em: {datetime.now().strftime("%d/%m/%Y às %H:%M")}\n'
                 'Documentação Técnica e Estratégica Completa').font.size = Pt(11)

    doc.add_page_break()

    # ── Índice ──
    _h1(doc, 'Sumário')
    toc = [
        '1.  Visão Geral do Projeto',
        '2.  Arquitetura do Sistema',
        '3.  Pipeline de Dados e Fluxo Geral',
        '4.  Coleta de Dados',
        '5.  Indicadores Técnicos',
        '6.  Motor de Direção',
        '7.  Motor de Volatilidade',
        '8.  Motor de Decisão Operacional',
        '9.  Estratégias de Opções BTC',
        '10. Interface Web',
        '11. Logger SQLite',
        '12. Telegram Bot',
        '13. Problemas Técnicos Enfrentados',
        '14. Evolução Futura',
        '15. Conclusão',
    ]
    for item in toc:
        doc.add_paragraph(item)
    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # Capítulo 1 — Visão Geral
    # ════════════════════════════════════════════════════
    _h1(doc, '1. Visão Geral do Projeto')

    _h2(doc, '1.1 Origem e Objetivo')
    _p(doc, 'O QSignalTrader nasceu da necessidade de um sistema quantitativo modular que integrasse '
            'análise de direção, volatilidade e decisão operacional em uma única plataforma. '
            'Diferente de robôs de execução automática, o foco está na geração de sinais analíticos '
            'para position trade — operações de médio e longo prazo em criptoativos e ações americanas.')
    _p(doc, 'O sistema foi projetado para ser uma ferramenta de apoio à decisão, não um executor automático '
            'de ordens. Ele oferece três motores independentes que convergem em uma leitura combinada: '
            'Direção (para onde o mercado tende), Volatilidade (qual o regime de movimento) e '
            'Decisão Operacional (o que fazer com essa informação).')

    _h2(doc, '1.2 Problema que Resolve')
    _bullet(doc, 'Análise dispersa: traders usam múltiplas ferramentas e indicadores desconectados.')
    _bullet(doc, 'Falta de contexto de volatilidade: muitos sistemas focam apenas na direção e ignoram o regime de volatilidade.')
    _bullet(doc, 'Dificuldade de integração: confluência entre múltiplos timeframes exige automação.')
    _bullet(doc, 'Opções de BTC requerem leitura combinada de direção E volatilidade.')

    _h2(doc, '1.3 Filosofia')
    _p(doc, 'O sistema parte de três premissas fundamentais:')
    _bullet(doc, 'Tendência é confirmada por alinhamento de EMAs e validação de RSI (RSI > 58 para alta, RSI < 42 para baixa).')
    _bullet(doc, 'Volatilidade não é direção. Um ativo pode ter alta volatilidade sem tendência definida.')
    _bullet(doc, 'Decisão operacional é a síntese de direção + volatilidade + contexto.')

    _h2(doc, '1.4 Ativos Suportados')
    _add_table(doc,
               ['Tipo', 'Ativos', 'Fonte de Dados', 'Timeframes'],
               [
                   ['Cripto (BTC)', 'BTC/USDT', 'CCXT (Binance)', '15m, 1h, 4h, 1D, 1W'],
                   ['Cripto (outras)', 'ETH/USDT, SOL/USDT, ...', 'CCXT (Binance)', '15m, 1h, 4h, 1D, 1W'],
                   ['Ações', 'AAPL, TSLA, NVDA, MSFT, SPY, QQQ', 'Yahoo Finance (yfinance)', '1d, 5d, 1wk'],
               ])
    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # Capítulo 2 — Arquitetura
    # ════════════════════════════════════════════════════
    _h1(doc, '2. Arquitetura do Sistema')

    _h2(doc, '2.1 Estrutura de Pastas')
    _code(doc, (
        'QSignalTrader/\n'
        '├── src/\n'
        '│   ├── data/\n'
        '│   │   ├── fetch_crypto.py      # Coleta OHLCV de cripto via CCXT\n'
        '│   │   └── yfinance_handler.py  # Coleta OHLCV de ações via yfinance\n'
        '│   ├── indicators/\n'
        '│   │   └── technicals.py        # EMAs, RSI, ATR, Fibonacci Pivots\n'
        '│   ├── strategy/\n'
        '│   │   └── multi_tf_analysis.py # Análise multi-timeframe (confluência)\n'
        '│   ├── signals/\n'
        '│   │   ├── signal_engine.py      # Motor de Direção + análise completa\n'
        '│   │   └── logger.py            # SQLite (sinais + volatilidade)\n'
        '│   ├── volatility/\n'
        '│   │   ├── volatility_config.py  # Constantes e funções auxiliares\n'
        '│   │   ├── realized_vol.py      # ATR%, BandWidth, RV7/30/90\n'
        '│   │   ├── implied_vol.py       # DVOL Deribit (BTC)\n'
        '│   │   ├── volatility_engine.py # Classificação de regime + scores\n'
        '│   │   └── options_strategy_engine.py  # Estratégias de opções BTC\n'
        '│   ├── decision/\n'
        '│   │   └── decision_engine.py   # Matriz direção × volatilidade\n'
        '│   └── utils/\n'
        '│       └── dates.py\n'
        '├── web/\n'
        '│   ├── app.py                   # FastAPI (endpoints)\n'
        '│   ├── static/\n'
        '│   │   ├── app.js               # Frontend interativo\n'
        '│   │   └── style.css            # Tema escuro\n'
        '│   └── templates/\n'
        '│       └── index.html           # Dashboard\n'
        '├── logs/\n'
        '│   └── sinais.db                # SQLite (criado automaticamente)\n'
        '├── abrir_qsignaltrader.bat      # Launcher Windows\n'
        '├── main.py\n'
        '├── requirements.txt\n'
        '└── README.md\n'
    ))

    _h2(doc, '2.2 Responsabilidades dos Módulos')
    _add_table(doc,
               ['Módulo', 'Responsabilidade'],
               [
                   ['data/fetch_crypto.py', 'Coleta OHLCV de criptomoedas via CCXT (Binance). Normaliza colunas.'],
                   ['data/yfinance_handler.py', 'Coleta OHLCV de ações via yfinance. Trata MultiIndex.'],
                   ['indicators/technicals.py', 'Cálculo de EMAs (8,21,50,100,200), RSI (14), ATR (14), Fibonacci pivots.'],
                   ['strategy/multi_tf_analysis.py', 'Análise de confluência entre 5 TFs (cripto) ou 3 TFs (ações).'],
                   ['signals/signal_engine.py', 'Motor de direção + orquestração dos 3 motores.'],
                   ['signals/logger.py', 'Persistência em SQLite: sinais e análises de volatilidade.'],
                   ['volatility/realized_vol.py', 'ATR%, Bollinger BandWidth, RV7/30/90.'],
                   ['volatility/implied_vol.py', 'DVOL via Deribit, IV Rank, IV Percentile (BTC).'],
                   ['volatility/volatility_engine.py', 'Classificação de regime de volatilidade + scores.'],
                   ['volatility/options_strategy_engine.py', 'Sugestão de estratégias de opções BTC.'],
                   ['decision/decision_engine.py', 'Matriz direção × volatilidade → decisão operacional.'],
                   ['web/app.py', 'FastAPI: /api/sinal, /api/analise-completa, /api/historico.'],
               ])

    _h2(doc, '2.3 Fluxo de Dados')
    _code(doc, (
        '  Símbolo (ex: BTC/USDT, AAPL)\n'
        '       │\n'
        '       ▼\n'
        '  ┌─────────────┐\n'
        '  │ Normalização │  → is_btc() + _normalizar_symbol()\n'
        '  └──────┬──────┘\n'
        '         │\n'
        '    ┌────┴────┐\n'
        '    ▼         ▼\n'
        '  Cripto     Ação\n'
        '  (CCXT)   (yfinance)\n'
        '    │         │\n'
        '    ▼         ▼\n'
        '  ┌─────────────────────┐\n'
        '  │ Indicadores Técnicos │  → EMAs, RSI, ATR\n'
        '  └─────────┬───────────┘\n'
        '            │\n'
        '       ┌────┴────┐\n'
        '       ▼         ▼\n'
        '  Direção    Volatilidade\n'
        '  (5 TFs)    (RV+IV p/ BTC)\n'
        '       │         │\n'
        '       └────┬────┘\n'
        '            ▼\n'
        '       Decisão\n'
        '    Operacional\n'
        '       │    \\\n'
        '       ▼     ▼\n'
        '  Frontend  SQLite\n'
        '  (Web)     (Log)'
    ))

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # Capítulo 3 — Pipeline
    # ════════════════════════════════════════════════════
    _h1(doc, '3. Pipeline de Dados e Fluxo Geral')

    _h2(doc, '3.1 Etapas do Pipeline')
    etapas = [
        ('1. Entrada', 'Usuário digita símbolo ou clica em botão rápido. Front-end chama GET /api/sinal?symbol=...'),
        ('2. Normalização', 'Símbolo é normalizado: BTCUSDT → BTC/USDT, BTC → BTC/USDT, AAPL permanece AAPL.'),
        ('3. Coleta OHLCV', 'Cripto usa CCXT/Binance. Ações usam yfinance. Dados convertidos para DataFrame pandas.'),
        ('4. Indicadores', 'EMAs 8/21/50/100/200 + RSI + ATR calculados em cada timeframe.'),
        ('5. Direção', 'Alinhamento de EMAs e RSI > 58 ou < 42 determinam tendência por timeframe. Confluência multi-TF gera classificação final.'),
        ('6. Volatilidade', 'Dados diários (1d) alimentam cálculo de ATR%, BandWidth, RV7/30/90. BTC também usa DVOL Deribit. Scores de expansão e contração determinam regime.'),
        ('7. Decisão', 'Matriz direção × volatilidade → decisão operacional (entrar comprado, short, aguardar, ficar fora).'),
        ('8. Opções BTC', 'Se is_btc=True, estratégias de opções são sugeridas via engine específica.'),
        ('9. Persistência', 'Resultado salvo em SQLite (tabelas sinais e volatilidade_btc).'),
        ('10. Saída', 'JSON é retornado ao front-end, que renderiza painéis de direção, volatilidade, decisão e opções.'),
    ]
    _add_table(doc, ['Etapa', 'Descrição'], etapas)

    _h2(doc, '3.2 Endpoints da API')
    _add_table(doc,
               ['Endpoint', 'Método', 'Descrição'],
               [
                   ['/api/sinal?symbol=...', 'GET', 'Análise completa (direção + volatilidade + decisão)'],
                   ['/api/analise-completa?symbol=...', 'GET', 'Mesmo que /api/sinal, endpoint explícito'],
                   ['/api/historico?limit=...', 'GET', 'Histórico de sinais do SQLite'],
                   ['/api/btc/volatilidade', 'GET', 'Apenas motor de volatilidade BTC'],
                   ['/api/btc/analise-completa', 'GET', 'Análise completa BTC (legado)'],
                   ['/api/btc/historico-volatilidade', 'GET', 'Histórico de análises de volatilidade BTC'],
               ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # Capítulo 4 — Coleta de Dados
    # ════════════════════════════════════════════════════
    _h1(doc, '4. Coleta de Dados')

    _h2(doc, '4.1 Criptomoedas — CCXT / Binance')
    _p(doc, 'A coleta de dados de criptomoedas utiliza a biblioteca CCXT, que abstrai mais de 100 exchanges '
            'em uma API unificada. Para o QSignalTrader, a exchange padrão é a Binance (spot market).')
    _p(doc, 'Função principal:')
    _code(doc, (
        'def get_crypto_data(symbol="BTC/USDT", exchange_name="binance",\n'
        '                    timeframe="1h", limit=100) -> pd.DataFrame:\n'
        '    exchange = ccxt.binance({"enableRateLimit": True})\n'
        '    ohlcv = exchange.fetch_ohlcv(symbol, timeframe.lower(), limit=limit)\n'
        '    df = pd.DataFrame(ohlcv, columns=["timestamp","open","high","low","close","volume"])\n'
        '    df["timestamp"] = pd.to_datetime(df["timestamp"], unit="ms")\n'
        '    df.set_index("timestamp", inplace=True)\n'
        '    return df'
    ))
    _p(doc, 'O parâmetro timeframe é convertido para minúsculo (timeframe.lower()) pois a API da Binance '
            'requer "1d" e "1w", não "1D" ou "1W". Sem essa conversão, a Binance retorna erro "Invalid interval".')

    _h2(doc, '4.2 Ações — Yahoo Finance / yfinance')
    _p(doc, 'Para ações americanas, utiliza-se a biblioteca yfinance, que acessa dados históricos do Yahoo Finance.')
    _p(doc, 'Classe principal:')
    _code(doc, (
        '@dataclass(frozen=True)\n'
        'class YahooFinanceDataHandler:\n'
        '    auto_adjust: bool = True\n'
        '\n'
        '    def fetch_ohlc(self, ticker, start=None, end=None,\n'
        '                   interval="1d", period=None) -> pd.DataFrame:\n'
        '        # Usa period ou start/end\n'
        '        df = yf.download(tickers=ticker, interval=interval,\n'
        '                         auto_adjust=self.auto_adjust, progress=False)\n'
        '        # Normaliza MultiIndex\n'
        '        if isinstance(df.columns, pd.MultiIndex):\n'
        '            df.columns = df.columns.get_level_values(0)\n'
        '        # Padroniza colunas: open, high, low, close, volume\n'
        '        return df[wanted].dropna()'
    ))
    _p(doc, 'Limitações do yfinance:')
    _bullet(doc, 'Não oferece dados intradiários gratuitos robustos (1m, 5m, 15m).')
    _bullet(doc, 'Os intervalos utilizados são 1d (diário) e 1wk (semanal).')
    _bullet(doc, 'Pode retornar MultiIndex nas colunas quando múltiplos tickers são baixados.')
    _bullet(doc, 'Dados de "Adj Close" podem ou não estar presentes dependendo de auto_adjust.')

    _h2(doc, '4.3 Normalização de Símbolos')
    _p(doc, 'O sistema aplica normalização robusta de símbolos para garantir consistência:')
    _code(doc, (
        'def _normalizar_symbol(symbol):\n'
        '    s = symbol.upper().strip()\n'
        '    if "/" not in s:\n'
        '        pairs = {"BTCUSDT": "BTC/USDT", "ETHUSDT": "ETH/USDT",\n'
        '                 "SOLUSDT": "SOL/USDT", ...}\n'
        '        if s in pairs: return pairs[s]\n'
        '    return s\n'
        '\n'
        'def is_btc(symbol):\n'
        '    s = symbol.upper().replace(" ", "")\n'
        '    return s in ("BTC", "BTCUSDT", "BTC/USDT", "BTC-USD")\n'
        '\n'
        'def normalizar_symbol_btc(symbol):\n'
        '    return "BTC/USDT" if is_btc(symbol) else symbol'
    ))

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # Capítulo 5 — Indicadores Técnicos
    # ════════════════════════════════════════════════════
    _h1(doc, '5. Indicadores Técnicos')
    _p(doc, 'Todos os indicadores são implementados no módulo indicators/technicals.py. '
            'Eles operam sobre DataFrames pandas com colunas padronizadas: open, high, low, close, volume.')

    _h2(doc, '5.1 EMAs — Médias Móveis Exponenciais')
    _p(doc, 'As EMAs são calculadas com pandas ewm(span=p, adjust=False):')
    _code(doc, (
        'def add_emas(df, periods=[8, 21, 50, 100, 200]):\n'
        '    out = df.copy()\n'
        '    for p in periods:\n'
        '        out[f"ema_{p}"] = out["close"].ewm(span=p, adjust=False).mean()\n'
        '    return out'
    ))
    _p(doc, 'O alinhamento das EMAs é o principal critério de tendência:')
    _bullet(doc, 'Tendência de alta: EMA 8 > EMA 21 > EMA 50 (cripto) ou EMA 8 > 21 > 50 > 100 > 200 (ações, baixa completa).')
    _bullet(doc, 'Tendência de baixa: EMA 8 < EMA 21 < EMA 50 < EMA 100 < EMA 200.')
    _bullet(doc, 'A EMA 200 não é exigida para alinhamento de alta (apenas 8, 21, 50), mas é usada para baixa.')

    _h2(doc, '5.2 RSI — Índice de Força Relativa')
    _p(doc, 'O RSI é calculado com período 14 usando o método de Wilder (média exponencial):')
    _code(doc, (
        'def add_rsi(df, period=14):\n'
        '    delta = df["close"].diff()\n'
        '    gain = delta.clip(lower=0)\n'
        '    loss = (-delta).clip(lower=0)\n'
        '    avg_gain = gain.ewm(alpha=1/period, adjust=False).mean()\n'
        '    avg_loss = loss.ewm(alpha=1/period, adjust=False).mean()\n'
        '    rs = avg_gain / avg_loss.replace(0, np.nan)\n'
        '    out["rsi"] = 100 - 100 / (1 + rs)\n'
        '    out["rsi"] = out["rsi"].fillna(50.0)\n'
        '    return out'
    ))
    _p(doc, 'Critérios de validação:')
    _add_table(doc,
               ['Condição', 'Valor', 'Interpretação'],
               [
                   ['RSI > 58', 'Força compradora', 'Tendência de alta é validada'],
                   ['RSI < 42', 'Força vendedora', 'Tendência de baixa é validada'],
                   ['42 ≤ RSI ≤ 58', 'Zona neutra', 'Sem validação direcional forte'],
               ])
    _p(doc, 'O valor de RSI > 58 foi definido empiricamente com base na observação de que o BTC '
            'tende a sustentar movimentos de alta apenas quando o RSI está acima desse patamar. '
            'O valor de RSI < 42 é o espelho simétrico para baixa. '
            'O sistema não usa RSI > 50 ou RSI < 50 como validação.')

    _h2(doc, '5.3 ATR — Average True Range')
    _p(doc, 'O ATR mede a volatilidade absoluta do preço:')
    _code(doc, (
        'def add_atr(df, period=14):\n'
        '    high, low, close = df["high"], df["low"], df["close"]\n'
        '    tr1 = high - low\n'
        '    tr2 = abs(high - close.shift(1))\n'
        '    tr3 = abs(low - close.shift(1))\n'
        '    true_range = pd.concat([tr1, tr2, tr3], axis=1).max(axis=1)\n'
        '    out["atr"] = true_range.ewm(alpha=1/period, adjust=False).mean()\n'
        '    return out'
    ))
    _p(doc, 'O ATR é usado de duas formas no sistema:'
            ' (1) como métrica absoluta na tabela de timeframes, '
            'e (2) como ATR percentual (ATR / close * 100) no motor de volatilidade.')

    _h2(doc, '5.4 Bollinger Bands — BandWidth')
    _p(doc, 'A largura das Bandas de Bollinger (BandWidth) é um indicador de compressão/expansão:')
    _code(doc, (
        'ma = close.rolling(20).mean()\n'
        'std = close.rolling(20).std()\n'
        'upper = ma + 2 * std\n'
        'lower = ma - 2 * std\n'
        'bandwidth = (upper - lower) / ma\n'
        'bandwidth_percentile = percentil na janela de 252 períodos'
    ))
    _p(doc, 'Um BandWidth percentil abaixo de 20% indica compressão (squeeze). '
            'Acima de 80% indica expansão. A inclinação do BandWidth (slope de 5 períodos) '
            'indica a direção da mudança.')

    _h2(doc, '5.5 Realized Volatility — RV')
    _p(doc, 'A volatilidade realizada é calculada usando retornos logarítmicos anualizados:')
    _code(doc, (
        'log_ret = ln(close / close.shift(1))\n'
        'rv_n = std(log_ret, janela=n) * sqrt(365)\n'
        '# n = 7, 30, 90 dias'
    ))
    _p(doc, 'Três janelas são analisadas:')
    _bullet(doc, 'RV7: volatilidade de curto prazo (7 dias).')
    _bullet(doc, 'RV30: volatilidade de médio prazo (30 dias).')
    _bullet(doc, 'RV90: volatilidade de longo prazo (90 dias).')
    _p(doc, 'As razões RV7/RV30 e RV30/RV90 indicam aceleração ou desaceleração da volatilidade:')
    _bullet(doc, 'RV7/RV30 > 1.25: volatilidade de curto prazo acelerando.')
    _bullet(doc, 'RV7/RV30 < 0.80: volatilidade de curto prazo esfriando.')

    _h2(doc, '5.6 Fibonacci Pivots')
    _p(doc, 'Os pivôs de Fibonacci são calculados a partir do OHLCV do período:')
    _code(doc, (
        'pivot = (high + low + close) / 3\n'
        'range_val = high - low\n'
        'r1 = pivot + range_val * 0.382\n'
        'r2 = pivot + range_val * 0.618\n'
        'r3 = pivot + range_val * 1.0\n'
        's1 = pivot - range_val * 0.382\n'
        's2 = pivot - range_val * 0.618\n'
        's3 = pivot - range_val * 1.0'
    ))
    _p(doc, 'Atualmente os pivôs são calculados mas não utilizados nos motores de decisão. '
            'Estão disponíveis como indicadores numéricos na interface web.')

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # Capítulo 6 — Motor de Direção
    # ════════════════════════════════════════════════════
    _h1(doc, '6. Motor de Direção')

    _h2(doc, '6.1 Conceito')
    _p(doc, 'O Motor de Direção determina o viés direcional do ativo com base em confluência '
            'de múltiplos timeframes. Cada timeframe é analisado individualmente e o conjunto '
            'dos resultados é classificado em regimes direcionais.')

    _h2(doc, '6.2 Análise por Timeframe')
    _p(doc, 'Para cada timeframe, a função _analisar_dataframe() calcula:')
    _code(doc, (
        'alinhamento_emas = EMA_8 > EMA_21 > EMA_50\n'
        'rsi_forte = RSI > 58\n'
        'tendencia_alta = alinhamento_emas AND rsi_forte\n'
        '\n'
        'alinhamento_emas_baixa = EMA_8 < EMA_21 < EMA_50 < EMA_100 < EMA_200\n'
        'rsi_fraco = RSI < 42\n'
        'tendencia_baixa = alinhamento_emas_baixa AND rsi_fraco'
    ))

    _h2(doc, '6.3 Timeframes por Tipo de Ativo')
    _add_table(doc,
               ['Ativo', 'Timeframes'],
               [
                   ['Cripto (BTC, ETH, SOL...)', '15m, 1h, 4h, 1D, 1W'],
                   ['Ações (AAPL, TSLA, SPY...)', '1d, 5d, 1wk'],
               ])

    _h2(doc, '6.4 Classificação de Regimes Direcionais')
    _p(doc, 'A função _classificar_direcao() mapeia o conjunto de timeframes em alta/baixa '
            'para um regime direcional. As combinações possíveis para cripto são:')

    _add_table(doc,
               ['Regime', 'Lado', 'Timeframes em Alta', 'Interpretação'],
               [
                   ['alta_forte', 'long', '1W, 1D, 4h, 1h, 15m', 'Todos os TFs alinhados em alta'],
                   ['alta_forte', 'long', '1D, 4h, 1h, 15m', 'Alta sem confirmação semanal'],
                   ['alta_moderada', 'long', '4h, 1h, 15m', 'Alta de curto prazo'],
                   ['alta_leve', 'long', '1h, 15m', 'Alta apenas intraday'],
                   ['possivel_reversao_alta', 'long', '1D', 'Apenas diário em alta'],
                   ['sem_tendencia_direcional', 'neutro', '15m isolado', 'Sinal muito fraco'],
                   ['transicao', 'neutro', '1D, 1h, 15m', 'Mistos entre TFs'],
               ])

    _p(doc, 'Para baixa, as combinações são simétricas (baixa_forte, baixa_moderada, etc.). '
            'Se não há timeframes em alta nem em baixa, o regime é sem_tendencia_direcional. '
            'Se há conflito entre alta e baixa, o regime é transicao.')

    _h2(doc, '6.5 Estrutura de Retorno')
    _code(doc, (
        '{\n'
        '    "regime_direcional": "alta_forte",\n'
        '    "interpretacao": "📈 Tendência de alta forte",\n'
        '    "lado": "long",\n'
        '    "forca": "forte",\n'
        '    "timeframes": { ... }  # análise individual de cada TF\n'
        '}'
    ))

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # Capítulo 7 — Motor de Volatilidade
    # ════════════════════════════════════════════════════
    _h1(doc, '7. Motor de Volatilidade')

    _h2(doc, '7.1 Conceito')
    _p(doc, 'O Motor de Volatilidade não tenta prever direção. Ele classifica o regime '
            'de volatilidade do ativo em uma de cinco categorias, baseado em métricas '
            'de volatilidade realizada (e implícita, apenas para BTC).')

    _h2(doc, '7.2 Fontes de Dados')
    _add_table(doc,
               ['Ativo', 'Vol. Realizada', 'Vol. Implícita'],
               [
                   ['BTC', 'ATR%, BandWidth, RV', 'DVOL (Deribit), IV Rank, IV/RV'],
                   ['Outras criptos', 'ATR%, BandWidth, RV', '—'],
                   ['Ações', 'ATR%, BandWidth, RV', '—'],
               ])

    _h2(doc, '7.3 Métricas de Volatilidade Realizada')
    _p(doc, 'O módulo realized_vol.py calcula:')
    _bullet(doc, 'ATR% = (ATR / close) × 100. Média móvel de 50 períodos e slope de 5 períodos.')
    _bullet(doc, 'Bollinger BandWidth = (banda_superior - banda_inferior) / média. Percentil em janela de 252. Slope de 5.')
    _bullet(doc, 'RV7, RV30, RV90 = desvio padrão dos retornos logarítmicos × √365. Razões RV7/RV30 e RV30/RV90.')

    _h2(doc, '7.4 DVOL — Volatilidade Implícita (BTC)')
    _p(doc, 'Para BTC, o sistema busca o DVOL (Deribit Volatility Index) via API pública da Deribit:')
    _code(doc, (
        'GET https://www.deribit.com/api/v2/public/get_volatility_index_data\n'
        '    ?currency=BTC&resolution=1D&start_timestamp=...&end_timestamp=...\n'
        '# Retorno: array de {timestamp, volatility} — ex: 55.0 = 55%'
    ))
    _p(doc, 'A partir do DVOL, são calculados:')
    _bullet(doc, 'DVOL atual, média 20, média 50, slope 3D e 7D.')
    _bullet(doc, 'IV Rank = (DVOL_atual - DVOL_min) / (DVOL_max - DVOL_min) × 100.')
    _bullet(doc, 'IV Percentile = % de dias históricos com DVOL abaixo do atual.')
    _bullet(doc, 'IV/RV Ratio = DVOL / (RV30 × 100). Interpretação: ≤1.00 = opções baratas; >1.25 = opções caras.')

    _h2(doc, '7.5 Scores de Expansão e Contração')
    _p(doc, 'Dois scores de 0 a 10 são calculados para quantificar a pressão de volatilidade:')

    _h3(doc, 'Score de Expansão')
    _p(doc, 'Para todos os ativos:')
    _bullet(doc, '+2 se ATR% > média ATR% 50 períodos.')
    _bullet(doc, '+2 se RV7/RV30 > 1.25.')
    _bullet(doc, '+2 se BandWidth está subindo após estar abaixo do percentil 20.')
    _p(doc, 'Adicional para BTC:')
    _bullet(doc, '+2 se DVOL > média 20 dias e DVOL subindo.')
    _bullet(doc, '+1 se IV Rank baixo/moderado e começou a subir.')
    _bullet(doc, '+1 se IV/RV ≤ 1.00.')
    _p(doc, 'Para não-BTC, o score bruto (máx 6) é normalizado para escala 0-10: '
            'score = round(bruto / 6 × 10).')

    _h3(doc, 'Score de Contração')
    _p(doc, 'Para todos os ativos:')
    _bullet(doc, '+1 se ATR% começou a cair.')
    _bullet(doc, '+1 se BandWidth começou a contrair.')
    _bullet(doc, '+2 se RV7/RV30 < 0.80.')
    _p(doc, 'Adicional para BTC:')
    _bullet(doc, '+2 se IV Rank > 70.')
    _bullet(doc, '+2 se IV Percentile > 70.')
    _bullet(doc, '+2 se DVOL caindo nos últimos 3-7 dias.')
    _bullet(doc, '+2 se IV/RV > 1.25.')

    _h2(doc, '7.6 Classificação de Regimes')
    _p(doc, 'O regime de volatilidade é classificado comparando os scores e métricas:')
    _add_table(doc,
               ['Regime', 'Ícone', 'Condição Principal'],
               [
                   ['Comprimida', '🧨', 'BandWidth %ile < 20, ATR% ratio < 1.0, RV ratio ≤ 1.0'],
                   ['Expansão', '🚀', 'ATR% ratio > 1.0, RV ratio > 1.0'],
                   ['Alta sustentada', '🌋', 'DVOL > 60, IV Rank > 60, ATR% ratio > 1.0'],
                   ['Contração provável', '🧊', 'IV Rank > 70, DVOL caindo'],
                   ['Transição', '🌫', 'Nenhuma condição anterior satisfeita'],
               ])
    _p(doc, 'Para ativos sem IV (não-BTC), a classificação usa apenas métricas realizadas, '
            'com fallbacks adaptados.')

    _h2(doc, '7.7 Confiança da Volatilidade')
    _add_table(doc,
               ['Ativo', 'Confiança', 'Motivo'],
               [
                   ['BTC', 'alta', 'Vol. realizada + implícita (DVOL)'],
                   ['Outras criptos', 'média', 'Apenas vol. realizada'],
                   ['Ações', 'média', 'Apenas vol. realizada'],
               ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # Capítulo 8 — Motor de Decisão
    # ════════════════════════════════════════════════════
    _h1(doc, '8. Motor de Decisão Operacional')

    _h2(doc, '8.1 Conceito')
    _p(doc, 'O Motor de Decisão combina os outputs do Motor de Direção e do Motor de Volatilidade '
            'para gerar uma recomendação operacional. Ele não substitui os outros motores — '
            'é a síntese deles.')

    _h2(doc, '8.2 Decisões Possíveis')
    _add_table(doc,
               ['Decisão', 'Ícone', 'Significado'],
               [
                   ['entrar_comprado', '🟢', 'Condições favoráveis para posição comprada'],
                   ['entrar_comprado_antecipado', '🟢', 'Entrada antes da confirmação total (vol. comprimida)'],
                   ['entrar_comprado_risco_definido', '🟡', 'Entrada possível, mas com risco controlado'],
                   ['entrar_vendido', '🔴', 'Condições favoráveis para posição vendida'],
                   ['entrar_vendido_antecipado', '🔴', 'Entrada short antes da confirmação total'],
                   ['entrar_vendido_risco_definido', '🟡', 'Entrada short com risco controlado'],
                   ['manter', '🟠', 'Manter posição, evitar nova entrada'],
                   ['entrar_com_confirmacao', '🟡', 'Aguardar confirmação antes de entrar'],
                   ['aguardar_rompimento', '🧨', 'Aguardar rompimento de range'],
                   ['ficar_de_fora', '⚪', 'Sem condições favoráveis para direcional'],
               ])

    _h2(doc, '8.3 Matriz de Decisão')
    _p(doc, 'A matriz abaixo mostra a decisão para cada combinação de direção e volatilidade. '
            'Ela contém aproximadamente 30 entradas cobrindo os cenários mais comuns.')

    _add_table(doc,
               ['Direção', 'Volatilidade', 'Decisão', 'Confiança'],
               [
                   ['alta_forte', 'compressao', '🟢 Entrada comprada antecipada', 'alta'],
                   ['alta_forte', 'expansao', '🟢 Entrada comprada favorecida', 'alta'],
                   ['alta_forte', 'alta_sustentada', '🟡 Entrada com risco definido', 'média'],
                   ['alta_forte', 'contracao', '🟠 Manter, evitar nova entrada', 'média'],
                   ['baixa_forte', 'compressao', '🔴 Entrada short antecipada', 'alta'],
                   ['baixa_forte', 'expansao', '🔴 Entrada short favorecida', 'alta'],
                   ['sem_tendencia', 'compressao', '🧨 Aguardar rompimento', 'média'],
                   ['sem_tendencia', 'expansao', '🟡 Aguardar direção ou confirmação', 'baixa'],
                   ['sem_tendencia', 'contracao', '⚪ Ficar de fora', 'alta'],
                   ['transicao', 'transicao', '⚪ Ficar de fora', 'alta'],
                   ['transicao', 'compressao', '🧨 Aguardar rompimento', 'média'],
               ])

    _h2(doc, '8.4 Estrutura de Retorno')
    _code(doc, (
        '{\n'
        '    "decisao": "🟢 Entrada comprada favorecida",\n'
        '    "decisao_key": "entrar_comprado",\n'
        '    "lado": "long",\n'
        '    "nivel": "entrada",\n'
        '    "confianca": "alta",\n'
        '    "explicacao": "Direção clara de alta com volatilidade em expansão...",\n'
        '    "alertas": ["Gerenciar risco com stop adequado", ...],\n'
        '    "opcoes_btc": null  # ou dict com estratégias, se is_btc\n'
        '}'
    ))

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # Capítulo 9 — Estratégias de Opções BTC
    # ════════════════════════════════════════════════════
    _h1(doc, '9. Estratégias de Opções BTC')

    _h2(doc, '9.1 Escopo')
    _p(doc, 'As estratégias de opções são exclusivas para BTC. Não são aplicadas a outras '
            'criptomoedas nem a ações. O sistema não executa ordens — apenas sugere '
            'estruturas de opções para estudo.')

    _h2(doc, '9.2 Motor de Estratégias')
    _p(doc, 'O módulo options_strategy_engine.py implementa a função sugerir_estrategias_btc() '
            'que recebe a direção, o regime de volatilidade, o score de expansão e de contração. '
            'A função mapeia a combinação para uma classificação de cenário e retorna:')
    _bullet(doc, 'Estratégias prioritárias.')
    _bullet(doc, 'Estratégias secundárias.')
    _bullet(doc, 'Estratégias a evitar ou usar com cuidado.')

    _h2(doc, '9.3 Exemplos de Mapeamento')
    _add_table(doc,
               ['Direção', 'Volatilidade', 'Estratégias Prioritárias', 'Evitar'],
               [
                   ['Alta forte', 'Expansão', 'Bull Call Spread, Long Call, Call Ratio Spread', 'Iron Condor, Short Strangle'],
                   ['Alta forte', 'Comprimida', 'Bull Call Spread, Long Call, Long Strangle', 'Iron Condor, Venda de vol.'],
                   ['Alta forte', 'Contração', 'Put Credit Spread, Call Debit Spread, Diagonal', 'Long Call, Straddle, Strangle'],
                   ['Baixa forte', 'Expansão', 'Bear Put Spread, Long Put, Long Strangle', 'Put Credit Spread, Iron Condor'],
                   ['Sem tendência', 'Comprimida', 'Long Straddle, Long Strangle, Calendar Spread', 'Call/Put direcional, Iron Condor'],
                   ['Sem tendência', 'Contração', 'Iron Condor, Short Strangle, Credit Spreads', 'Straddle, Strangle, Compra de opções'],
               ])

    _h2(doc, '9.4 Estratégias Nomeadas')
    _p(doc, 'As seguintes estratégias são referenciadas pelo motor:')
    _add_table(doc,
               ['Estratégia', 'Tipo', 'Descrição'],
               [
                   ['Bull Call Spread', 'Direcional altista', 'Compra de call ITM + venda de call OTM'],
                   ['Bear Put Spread', 'Direcional baixista', 'Compra de put ITM + venda de put OTM'],
                   ['Put Credit Spread', 'Venda de vol. altista', 'Venda de put OTM + compra de put mais OTM'],
                   ['Long Straddle', 'Neutro comprado', 'Compra de call ATM + put ATM'],
                   ['Long Strangle', 'Neutro comprado', 'Compra de call OTM + put OTM'],
                   ['Iron Condor', 'Neutro vendido', 'Venda de call spread + put spread OTM'],
                   ['Calendar Spread', 'Tempo/vol.', 'Venda de opção curta + compra de opção longa'],
                   ['Diagonal Spread', 'Direcional + tempo', 'Venda curta + compra longa em strikes diferentes'],
               ])

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # Capítulo 10 — Interface Web
    # ════════════════════════════════════════════════════
    _h1(doc, '10. Interface Web')

    _h2(doc, '10.1 Tecnologias')
    _p(doc, 'A interface web utiliza:')
    _bullet(doc, 'FastAPI — framework Python para API REST.')
    _bullet(doc, 'Jinja2 — template engine para HTML.')
    _bullet(doc, 'HTML/CSS/JS puro — frontend leve, sem frameworks pesados.')
    _bullet(doc, 'Tema escuro — design moderno com fundo #0d1117 (GitHub Dark).')

    _h2(doc, '10.2 Painéis')
    _p(doc, 'O dashboard exibe os seguintes blocos:')
    _add_table(doc,
               ['Painel', 'Visível para', 'Conteúdo'],
               [
                   ['Direção', 'Todos', 'Símbolo, tipo, interpretação, timestamp'],
                   ['Timeframes', 'Todos', 'Tabela com EMAs, RSI, ATR, tendência por TF'],
                   ['Volatilidade', 'Todos', 'Regime, scores, métricas realizadas, leitura'],
                   ['Decisão', 'Todos', 'Decisão, lado, confiança, explicação, alertas'],
                   ['Opções BTC', 'BTC apenas', 'Estratégias prioritárias, secundárias, evitar'],
                   ['Histórico', 'Todos', 'Sinais anteriores do SQLite'],
               ])

    _h2(doc, '10.3 Cores e Significados')
    _add_table(doc,
               ['Cor', 'Contexto', 'Significado'],
               [
                   ['Verde', 'Direção/Decisão', 'Alta / entrada comprada'],
                   ['Vermelho', 'Direção/Decisão', 'Baixa / entrada vendida'],
                   ['Amarelo', 'Decisão', 'Confirmar / aguardar / risco definido'],
                   ['Laranja', 'Decisão', 'Manter / reduzir / aguardar rompimento'],
                   ['Cinza', 'Tendência/Decisão', 'Sem tendência / ficar de fora'],
               ])

    _h2(doc, '10.4 Endpoints')
    _p(doc, 'O backend FastAPI oferece os seguintes endpoints REST:')
    _code(doc, (
        'GET  /                               → Dashboard HTML\n'
        'GET  /api/sinal?symbol=BTC/USDT       → Análise completa (JSON)\n'
        'GET  /api/analise-completa?symbol=... → Análise completa explícita\n'
        'GET  /api/historico?limit=25          → Histórico de sinais\n'
        'GET  /api/btc/volatilidade            → Volatilidade BTC\n'
        'GET  /api/btc/analise-completa        → Análise BTC completa\n'
        'GET  /api/btc/historico-volatilidade   → Histórico de vol. BTC'
    ))

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # Capítulo 11 — SQLite Logger
    # ════════════════════════════════════════════════════
    _h1(doc, '11. Logger SQLite')

    _h2(doc, '11.1 Estrutura do Banco')
    _p(doc, 'O banco de dados SQLite está localizado em logs/sinais.db e contém duas tabelas.')

    _h3(doc, 'Tabela: sinais')
    _code(doc, (
        'CREATE TABLE sinais (\n'
        '    id INTEGER PRIMARY KEY AUTOINCREMENT,\n'
        '    timestamp TEXT NOT NULL,\n'
        '    ativo TEXT NOT NULL,\n'
        '    tipo TEXT NOT NULL,\n'
        '    interpretacao TEXT NOT NULL\n'
        ');'
    ))
    _p(doc, 'Registra cada sinal de direção gerado para qualquer ativo.')

    _h3(doc, 'Tabela: volatilidade_btc')
    _code(doc, (
        'CREATE TABLE volatilidade_btc (\n'
        '    id INTEGER PRIMARY KEY AUTOINCREMENT,\n'
        '    timestamp TEXT, symbol TEXT, regime TEXT,\n'
        '    score_expansao INTEGER, score_contracao INTEGER,\n'
        '    atr_percent REAL, bandwidth_percentile REAL,\n'
        '    rv7 REAL, rv30 REAL, rv90 REAL,\n'
        '    dvol REAL, iv_rank REAL, iv_percentile REAL,\n'
        '    iv_rv_ratio REAL, leitura TEXT\n'
        ');'
    ))
    _p(doc, 'Registra as análises de volatilidade BTC com todas as métricas.')

    _h2(doc, '11.2 Funções de Acesso')
    _bullet(doc, 'salvar_sinal(ativo, tipo, interpretacao) — insere na tabela sinais.')
    _bullet(doc, 'carregar_historico(limit, ativo, tipo) — consulta com filtros.')
    _bullet(doc, 'salvar_volatilidade_btc(vol_data) — insere na tabela volatilidade_btc.')
    _bullet(doc, 'carregar_historico_volatilidade(limit) — consulta com limite.')

    _h2(doc, '11.3 Objetivo Futuro')
    _p(doc, 'O histórico em SQLite foi projetado para permitir futuros backtests e análises '
            'de desempenho dos sinais ao longo do tempo.')

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # Capítulo 12 — Telegram Bot
    # ════════════════════════════════════════════════════
    _h1(doc, '12. Telegram Bot')
    _p(doc, 'O sistema possui integração com Telegram Bot para envio de sinais. '
            'O bot permite consultar sinais sob demanda via comando /sinal '
            'e possui agendamento automático em horários definidos.')
    _p(doc, 'Funcionalidades:')
    _bullet(doc, 'Comando /sinal — retorna sinal do BTC.')
    _bullet(doc, 'Acesso restrito por ID de usuário autorizado.')
    _bullet(doc, 'Timezone corrigido com pytz (America/Sao_Paulo).')
    _bullet(doc, 'Agendamentos em horários personalizados.')
    _p(doc, 'Nota: o bot Telegram não foi o foco desta iteração de desenvolvimento '
            'e pode exigir ajustes para funcionar com a nova estrutura de 3 motores.')

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # Capítulo 13 — Problemas Técnicos
    # ════════════════════════════════════════════════════
    _h1(doc, '13. Problemas Técnicos Enfrentados')

    problemas = [
        ('Importações lentas (pandas ~30s)',
         'O pandas demora ~30 segundos para importar pela primeira vez no Windows, fazendo '
         'com que o primeiro request HTTP ao servidor demore muito.',
         'O timeout do frontend foi aumentado de 60s para 300s (específico para BTC que faz mais chamadas). '
         'A resposta é: aceitar a latência da primeira chamada.'),
        ('numpy.bool_ não serializável para JSON',
         'Os valores booleanos retornados pelas funções de análise eram numpy.bool_, que não são '
         'serializáveis pelo jsonable_encoder do FastAPI, causando HTTP 500 com corpo vazio.',
         'Envolver todos os valores booleanos e númericos com bool() e float() explícitos. '
         'Criar funções _safe_float() e _safe_scalar() em multi_tf_analysis.py.'),
        ('Symbol com "/" quebrando URL do FastAPI',
         'O símbolo BTC/USDT contém "/", que ao ser usado como path parameter '
         '(/api/sinal/{symbol}) quebrava a rota.',
         'Trocar path parameter para query parameter: /api/sinal?symbol=BTC/USDT. '
         'Adicionar encodeURIComponent() no front-end.'),
        ('classList null no frontend',
         'Ao remover o card "leitura combinada", o JavaScript ainda referenciava '
         'document.getElementById("leituraCombinadaCard"), que não existia mais.',
         'Remover o bloco de código órfão que referenciava o ID removido.'),
        ('Divergência de timeframes entre cripto e ações',
         'O yfinance não suporta timeframes intradiários confiáveis, enquanto o ccxt suporta.',
         'Manter timeframes separados por tipo de ativo: 15m/1h/4h/1D/1W para cripto, '
         '1d/5d/1wk para ações.'),
        ('Intervalo inválido na Binance (1D vs 1d)',
         'A API da Binance rejeita intervalos maiúsculos como "1D" e "1W", aceitando apenas '
         '"1d" e "1w". O ccxt não normaliza o case automaticamente.',
         'Converter o timeframe para minúsculo com .lower() antes de chamar fetch_ohlcv().'),
        ('Dados semanais insuficientes para indicadores',
         'Com limit=200 e timeframe semanal, a EMA 50 não converge adequadamente se houver '
         'menos dados que o esperado.',
         'Aumentar limit para 500 em timeframes diário e semanal. Para ações, usar period="max" '
         'no yfinance para obter o histórico completo.'),
        ('Stocks não recebiam volatilidade — só BTC',
         'O motor de volatilidade original só funcionava para BTC. Ações e outras criptos '
         'não tinham painel de volatilidade.',
         'Generalizar a função analisar_volatilidade() para aceitar qualquer símbolo. '
         'Para não-BTC, usar apenas métricas realizadas com scores normalizados.'),
    ]
    for titulo, desc, solucao in problemas:
        _h3(doc, titulo)
        _p(doc, desc)
        _p(doc, f'Solução: {solucao}')

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # Capítulo 14 — Evolução Futura
    # ════════════════════════════════════════════════════
    _h1(doc, '14. Evolução Futura')

    futuros = [
        ('Machine Learning para Regimes',
         'Treinar modelos para classificar automaticamente os regimes de volatilidade '
         'e direção, reduzindo a dependência de regras fixas.'),
        ('Backtesting Robusto',
         'Implementar um motor de backtesting que avalie o desempenho histórico das '
         'decisões operacionais usando o banco SQLite existente.'),
        ('Options Flow / Tape Reading',
         'Integrar dados de fluxo de opções (put/call ratio, open interest, delta) '
         'para refinar as sugestões de estratégias BTC.'),
        ('IV Rank e Skew Reais',
         'Expandir a coleta de dados de opções para obter IV real por vencimento e '
         'skew (diferença entre put e call IV), em vez de apenas DVOL.'),
        ('Integração Deribit Completa',
         'Conectar à API da Deribit para dados de opções em tempo real, incluindo '
         'Greek values, superfície de volatilidade e estrutura a termo.'),
        ('Risk Engine',
         'Criar um motor de risco que calcule Value-at-Risk, exposição a Greek e '
         'tamanho de posição sugerido com base na volatilidade e capital disponível.'),
        ('Portfolio Engine',
         'Analisar correlação entre múltiplos ativos e sugerir alocações que '
         'considerem o regime de volatilidade de cada um.'),
        ('Scanner Multi-Ativo',
         'Escaneiar automaticamente centenas de ativos e destacar aqueles com '
         'configurações interessantes de direção + volatilidade.'),
        ('Alertas Inteligentes',
         'Notificar o usuário quando houver mudança significativa no regime de '
         'volatilidade ou na direção de ativos monitorados.'),
    ]
    for titulo, desc in futuros:
        _h3(doc, titulo)
        _p(doc, desc)

    doc.add_page_break()

    # ════════════════════════════════════════════════════
    # Capítulo 15 — Conclusão
    # ════════════════════════════════════════════════════
    _h1(doc, '15. Conclusão')

    _h2(doc, '15.1 Filosofia do Sistema')
    _p(doc, 'O QSignalTrader foi construído sobre três pilares: direção, volatilidade e decisão. '
            'A separação entre eles é deliberada — cada motor opera de forma independente, '
            'com métricas próprias e lógica específica. A convergência ocorre apenas no '
            'momento da decisão operacional, quando as leituras de todos os motores são '
            'combinadas em uma síntese coerente.')

    _h2(doc, '15.2 Importância do Contexto')
    _p(doc, 'Um dos aprendizados centrais do projeto é que direção sem contexto de volatilidade '
            'é insuficiente para a tomada de decisão. Um ativo pode estar em tendência de alta '
            'forte, mas se a volatilidade já está excessivamente elevada, a relação '
            'risco/retorno para uma nova entrada pode ser desfavorável. Da mesma forma, '
            'volatilidade sem direção pode indicar movimento caótico, não oportunidade.')

    _h2(doc, '15.3 União entre Direção e Volatilidade')
    _p(doc, 'A grande inovação do sistema é tratar direção e volatilidade como variáveis '
            'independentes que se cruzam. Isso permite:')
    _bullet(doc, 'Identificar entradas favoráveis (alta + expansão de vol).')
    _bullet(doc, 'Evitar entradas tardias (alta + contração de vol).')
    _bullet(doc, 'Preparar-se para movimentos (compressão + direção indefinida).')
    _bullet(doc, 'Sugerir estratégias de opções específicas para BTC (alta + expansão → Bull Call Spread; '
                 'sem tendência + compressão → Long Straddle).')

    _h2(doc, '15.4 Visão Futura')
    _p(doc, 'O QSignalTrader é uma plataforma em evolução. A arquitetura modular permite que '
            'novos motores sejam adicionados sem reescrever os existentes. Machine learning, '
            'backtesting quantitativo, análise de fluxo de opções e integração com exchanges '
            'são caminhos naturais de evolução. O importante é que a base — a separação '
            'entre direção, volatilidade e decisão — permaneça como fundamento.')

    _h2(doc, '15.5 Dependências do Projeto')
    _add_table(doc,
               ['Biblioteca', 'Uso'],
               [
                   ['yfinance', 'Coleta de dados de ações'],
                   ['pandas', 'Manipulação de dados tabulares'],
                   ['numpy', 'Operações matemáticas'],
                   ['ccxt', 'Coleta de dados de criptomoedas'],
                   ['fastapi', 'API REST'],
                   ['uvicorn', 'Servidor ASGI'],
                   ['jinja2', 'Templates HTML'],
                   ['pytz', 'Timezones'],
                   ['requests', 'Chamadas HTTP (DVOL Deribit)'],
                   ['python-docx', 'Geração de documentação (este documento)'],
               ])

    # ── Salvar ──
    output = r'C:\Users\angel\OneDrive\Documentos\Documentos\QSignalTrader\QSignalTrader_Documentacao_Completa.docx'
    doc.save(output)
    print('OK - Documentacao gerada.')
    print(f'Arquivo: {output}')
    return output


if __name__ == '__main__':
    criar_documentacao()
