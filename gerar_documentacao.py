"""
Script para gerar documentação do Projeto QT em formato DOCX
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def adicionar_titulo(doc, texto, nivel=1):
    """Adiciona um título formatado ao documento"""
    heading = doc.add_heading(texto, level=nivel)
    return heading

def adicionar_paragrafo(doc, texto, negrito=False):
    """Adiciona um parágrafo ao documento"""
    p = doc.add_paragraph(texto)
    if negrito:
        p.runs[0].bold = True
    return p

def adicionar_codigo(doc, codigo, linguagem="Python"):
    """Adiciona um bloco de código formatado"""
    p = doc.add_paragraph(codigo)
    # Formatar como código
    run = p.runs[0]
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    # Adicionar borda/fundo cinza se desejar
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), 'F0F0F0')
    p._element.get_or_add_pPr().append(shd)
    
    return p

def criar_documentacao():
    """Cria o documento de documentação do projeto QT"""
    
    # Criar documento
    doc = Document()
    
    # Título principal
    titulo = doc.add_heading('Documentação do Projeto QT', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtítulo com data
    subtitulo = doc.add_paragraph(f'Gerado em: {datetime.now().strftime("%d/%m/%Y às %H:%M")}')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # ==================== ÍNDICE ====================
    adicionar_titulo(doc, 'Índice', 1)
    doc.add_paragraph('1. Visão Geral do Projeto')
    doc.add_paragraph('2. Estrutura do Projeto')
    doc.add_paragraph('3. Dependências')
    doc.add_paragraph('4. Módulos Implementados')
    doc.add_paragraph('5. Como Usar')
    doc.add_paragraph('6. Testes')
    doc.add_paragraph('7. Próximos Passos')
    doc.add_paragraph('8. Manutenção')
    
    doc.add_page_break()
    
    # ==================== 1. VISÃO GERAL ====================
    adicionar_titulo(doc, '1. Visão Geral do Projeto', 1)
    
    doc.add_paragraph(
        'O Projeto QT é uma aplicação Python focada em análise quantitativa de mercado financeiro. '
        'O objetivo principal é fornecer ferramentas para coleta, processamento e análise de dados '
        'financeiros de ações, com foco inicial na integração com Yahoo Finance.'
    )
    
    adicionar_titulo(doc, 'Objetivos', 2)
    doc.add_paragraph('✓ Coletar dados históricos de ações via API do Yahoo Finance', style='List Bullet')
    doc.add_paragraph('✓ Processar e normalizar dados para análise quantitativa', style='List Bullet')
    doc.add_paragraph('✓ Fornecer base para implementação de estratégias de trading', style='List Bullet')
    doc.add_paragraph('✓ Manter código modular e testável', style='List Bullet')
    
    # ==================== 2. ESTRUTURA ====================
    adicionar_titulo(doc, '2. Estrutura do Projeto', 1)
    
    doc.add_paragraph(
        'O projeto segue uma estrutura modular organizada da seguinte forma:'
    )
    
    estrutura = """
QT/
├── src/
│   ├── data/
│   │   └── yfinance_handler.py    # Handler para coleta de dados do Yahoo Finance
│   └── utils/
│       └── dates.py                # Utilitários de datas (placeholder)
├── tests/
│   └── test_datahandler.py        # Testes unitários (placeholder)
├── .venv/                          # Ambiente virtual Python
├── requirements.txt                # Dependências do projeto
└── test_quick.py                   # Script de teste rápido
"""
    
    p = doc.add_paragraph(estrutura)
    p_format = p.paragraph_format
    p_format.left_indent = Inches(0.5)
    
    adicionar_titulo(doc, 'Descrição dos Diretórios', 2)
    
    doc.add_paragraph('src/data/: Módulos de coleta e manipulação de dados financeiros', style='List Bullet')
    doc.add_paragraph('src/utils/: Funções utilitárias compartilhadas', style='List Bullet')
    doc.add_paragraph('tests/: Testes unitários e de integração', style='List Bullet')
    doc.add_paragraph('.venv/: Ambiente virtual isolado para dependências', style='List Bullet')
    
    # ==================== 3. DEPENDÊNCIAS ====================
    adicionar_titulo(doc, '3. Dependências', 1)
    
    doc.add_paragraph(
        'O projeto utiliza as seguintes bibliotecas Python:'
    )
    
    # Tabela de dependências
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Light Grid Accent 1'
    
    # Cabeçalho
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Biblioteca'
    hdr_cells[1].text = 'Descrição'
    
    # Linhas
    table.rows[1].cells[0].text = 'yfinance'
    table.rows[1].cells[1].text = 'Biblioteca para download de dados do Yahoo Finance'
    
    table.rows[2].cells[0].text = 'pandas'
    table.rows[2].cells[1].text = 'Manipulação e análise de dados estruturados'
    
    table.rows[3].cells[0].text = 'numpy'
    table.rows[3].cells[1].text = 'Operações matemáticas e arrays multidimensionais'
    
    doc.add_paragraph()
    
    adicionar_titulo(doc, 'Instalação de Dependências', 2)
    doc.add_paragraph('Para instalar todas as dependências necessárias, execute:')
    adicionar_codigo(doc, 'pip install -r requirements.txt')
    
    # ==================== 4. MÓDULOS IMPLEMENTADOS ====================
    adicionar_titulo(doc, '4. Módulos Implementados', 1)
    
    adicionar_titulo(doc, '4.1 YahooFinanceDataHandler', 2)
    
    doc.add_paragraph('Localização: src/data/yfinance_handler.py', style='Intense Quote')
    
    doc.add_paragraph(
        'Este é o módulo principal para coleta de dados financeiros. Implementa uma classe '
        'que encapsula a funcionalidade do yfinance e padroniza os dados retornados.'
    )
    
    adicionar_titulo(doc, 'Características', 3)
    doc.add_paragraph('Implementado como dataclass frozen (imutável)', style='List Bullet')
    doc.add_paragraph('Suporte para ajuste automático de preços', style='List Bullet')
    doc.add_paragraph('Padronização de nomes de colunas (minúsculas)', style='List Bullet')
    doc.add_paragraph('Validação de dados retornados', style='List Bullet')
    doc.add_paragraph('Remoção automática de valores NaN', style='List Bullet')
    
    adicionar_titulo(doc, 'Método Principal: fetch_ohlc()', 3)
    
    doc.add_paragraph('Parâmetros:')
    param_table = doc.add_table(rows=5, cols=3)
    param_table.style = 'Light List Accent 1'
    
    param_table.rows[0].cells[0].text = 'Parâmetro'
    param_table.rows[0].cells[1].text = 'Tipo'
    param_table.rows[0].cells[2].text = 'Descrição'
    
    param_table.rows[1].cells[0].text = 'ticker'
    param_table.rows[1].cells[1].text = 'str'
    param_table.rows[1].cells[2].text = 'Símbolo da ação (ex: AAPL, PETR4.SA)'
    
    param_table.rows[2].cells[0].text = 'start'
    param_table.rows[2].cells[1].text = 'str'
    param_table.rows[2].cells[2].text = 'Data de início (formato: YYYY-MM-DD)'
    
    param_table.rows[3].cells[0].text = 'end'
    param_table.rows[3].cells[1].text = 'Optional[str]'
    param_table.rows[3].cells[2].text = 'Data de fim (opcional, padrão: hoje)'
    
    param_table.rows[4].cells[0].text = 'interval'
    param_table.rows[4].cells[1].text = 'str'
    param_table.rows[4].cells[2].text = 'Intervalo dos dados (padrão: 1d)'
    
    doc.add_paragraph()
    doc.add_paragraph('Retorno: DataFrame pandas com colunas: open, high, low, close, volume')
    
    adicionar_titulo(doc, 'Exemplo de Uso', 3)
    
    codigo_exemplo = """from src.data.yfinance_handler import YahooFinanceDataHandler

# Criar instância do handler
handler = YahooFinanceDataHandler(auto_adjust=True)

# Baixar dados
df = handler.fetch_ohlc(
    ticker="AAPL",
    start="2023-01-01",
    end="2023-03-01"
)

# Visualizar dados
print(df.head())"""
    
    adicionar_codigo(doc, codigo_exemplo)
    
    adicionar_titulo(doc, '4.2 Módulo dates.py', 2)
    doc.add_paragraph('Localização: src/utils/dates.py', style='Intense Quote')
    doc.add_paragraph(
        'Atualmente é um arquivo placeholder. Futuramente conterá funções utilitárias '
        'para manipulação de datas, como conversão de formatos, validação, cálculo de '
        'dias úteis, etc.'
    )
    
    # ==================== 5. COMO USAR ====================
    adicionar_titulo(doc, '5. Como Usar o Projeto', 1)
    
    adicionar_titulo(doc, 'Passo 1: Configurar Ambiente', 2)
    
    setup_code = """# Criar ambiente virtual
python -m venv .venv

# Ativar ambiente virtual (Windows)
.venv\\Scripts\\activate

# Instalar dependências
pip install -r requirements.txt"""
    
    adicionar_codigo(doc, setup_code)
    
    adicionar_titulo(doc, 'Passo 2: Executar Teste Rápido', 2)
    doc.add_paragraph('O arquivo test_quick.py demonstra o uso básico do sistema:')
    
    adicionar_codigo(doc, 'python test_quick.py')
    
    doc.add_paragraph(
        'Este script baixa dados da Apple (AAPL) de janeiro a março de 2023 e exibe '
        'as primeiras e últimas linhas, além de informações sobre as colunas.'
    )
    
    adicionar_titulo(doc, 'Passo 3: Integrar em Seus Scripts', 2)
    doc.add_paragraph('Para usar em seus próprios scripts:')
    
    integracao_code = """from src.data.yfinance_handler import YahooFinanceDataHandler

# Criar handler
dh = YahooFinanceDataHandler(auto_adjust=True)

# Baixar dados de múltiplos períodos
df_curto = dh.fetch_ohlc("MSFT", start="2024-01-01", end="2024-03-01")
df_longo = dh.fetch_ohlc("GOOGL", start="2020-01-01", end="2024-01-01")

# Processar dados conforme necessário
print(f"Dados MSFT: {len(df_curto)} registros")
print(f"Dados GOOGL: {len(df_longo)} registros")"""
    
    adicionar_codigo(doc, integracao_code)
    
    # ==================== 6. TESTES ====================
    adicionar_titulo(doc, '6. Testes', 1)
    
    doc.add_paragraph(
        'O projeto possui estrutura preparada para testes unitários na pasta tests/. '
        'Atualmente, test_datahandler.py é um placeholder.'
    )
    
    adicionar_titulo(doc, 'Plano de Testes Futuros', 2)
    doc.add_paragraph('Testes unitários para YahooFinanceDataHandler', style='List Bullet')
    doc.add_paragraph('Testes de validação de dados retornados', style='List Bullet')
    doc.add_paragraph('Testes de tratamento de erros (ticker inválido, período sem dados)', style='List Bullet')
    doc.add_paragraph('Testes de integração com yfinance', style='List Bullet')
    doc.add_paragraph('Testes de performance para grandes volumes de dados', style='List Bullet')
    
    adicionar_titulo(doc, 'Framework Recomendado', 2)
    doc.add_paragraph('pytest: Framework de testes moderno e poderoso')
    doc.add_paragraph('pytest-cov: Para análise de cobertura de código')
    
    # ==================== 7. PRÓXIMOS PASSOS ====================
    adicionar_titulo(doc, '7. Próximos Passos', 1)
    
    doc.add_paragraph(
        'Sugestões de funcionalidades e melhorias para o desenvolvimento futuro do projeto:'
    )
    
    adicionar_titulo(doc, '7.1 Curto Prazo', 2)
    doc.add_paragraph('Implementar funções em dates.py para manipulação de datas', style='List Bullet')
    doc.add_paragraph('Criar testes unitários completos', style='List Bullet')
    doc.add_paragraph('Adicionar logging para debug e monitoramento', style='List Bullet')
    doc.add_paragraph('Implementar cache de dados para evitar downloads repetidos', style='List Bullet')
    doc.add_paragraph('Documentar código com docstrings completas', style='List Bullet')
    
    adicionar_titulo(doc, '7.2 Médio Prazo', 2)
    doc.add_paragraph('Criar módulo de indicadores técnicos (RSI, MACD, Bollinger Bands)', style='List Bullet')
    doc.add_paragraph('Implementar sistema de estratégias de trading', style='List Bullet')
    doc.add_paragraph('Adicionar suporte para múltiplas fontes de dados', style='List Bullet')
    doc.add_paragraph('Criar visualizações de dados com matplotlib/plotly', style='List Bullet')
    doc.add_paragraph('Implementar backtesting de estratégias', style='List Bullet')
    
    adicionar_titulo(doc, '7.3 Longo Prazo', 2)
    doc.add_paragraph('Desenvolver interface web para visualização', style='List Bullet')
    doc.add_paragraph('Implementar machine learning para previsões', style='List Bullet')
    doc.add_paragraph('Criar API REST para acesso aos dados', style='List Bullet')
    doc.add_paragraph('Adicionar alertas em tempo real', style='List Bullet')
    doc.add_paragraph('Integração com corretoras para trading automático', style='List Bullet')
    
    # ==================== 8. MANUTENÇÃO ====================
    adicionar_titulo(doc, '8. Guia de Manutenção', 1)
    
    adicionar_titulo(doc, '8.1 Estrutura de Código', 2)
    doc.add_paragraph(
        'O projeto segue princípios de código limpo e arquitetura modular:'
    )
    doc.add_paragraph('Dataclasses: Uso de @dataclass para estruturas de dados imutáveis', style='List Bullet')
    doc.add_paragraph('Type Hints: Tipagem estática para melhor manutenibilidade', style='List Bullet')
    doc.add_paragraph('Separação de Responsabilidades: Cada módulo tem uma função específica', style='List Bullet')
    doc.add_paragraph('Padronização: Convenções consistentes (nomes em minúsculas, etc.)', style='List Bullet')
    
    adicionar_titulo(doc, '8.2 Convenções de Código', 2)
    doc.add_paragraph('PEP 8: Seguir guia de estilo oficial do Python')
    doc.add_paragraph('Docstrings: Documentar todas as classes e funções públicas')
    doc.add_paragraph('Type Hints: Usar em todas as assinaturas de funções')
    doc.add_paragraph('Imports: Organizar em ordem alfabética e por tipo')
    
    adicionar_titulo(doc, '8.3 Adicionando Novos Módulos', 2)
    doc.add_paragraph(
        'Ao adicionar novos módulos ao projeto, siga estas diretrizes:'
    )
    
    diretrizes_code = """# 1. Criar arquivo no diretório apropriado
# src/data/ - para handlers de dados
# src/utils/ - para utilitários gerais
# src/strategies/ - para estratégias de trading (futuro)

# 2. Estrutura básica de um novo módulo
from __future__ import annotations
from typing import Optional
import pandas as pd

class NovoModulo:
    \"\"\"Descrição clara do propósito do módulo.\"\"\"
    
    def __init__(self, param: str):
        \"\"\"
        Inicializar o módulo.
        
        Args:
            param: Descrição do parâmetro
        \"\"\"
        self.param = param
    
    def metodo_principal(self, data: pd.DataFrame) -> pd.DataFrame:
        \"\"\"
        Descrição do que o método faz.
        
        Args:
            data: DataFrame de entrada
            
        Returns:
            DataFrame processado
        \"\"\"
        # Implementação
        return data

# 3. Criar testes correspondentes em tests/"""
    
    adicionar_codigo(doc, diretrizes_code)
    
    adicionar_titulo(doc, '8.4 Atualizando Dependências', 2)
    doc.add_paragraph('Verificar versões atuais:')
    adicionar_codigo(doc, 'pip list --outdated')
    
    doc.add_paragraph()
    doc.add_paragraph('Atualizar requirements.txt:')
    adicionar_codigo(doc, 'pip freeze > requirements.txt')
    
    adicionar_titulo(doc, '8.5 Resolução de Problemas Comuns', 2)
    
    # Tabela de problemas
    prob_table = doc.add_table(rows=4, cols=2)
    prob_table.style = 'Medium Grid 1 Accent 1'
    
    prob_table.rows[0].cells[0].text = 'Problema'
    prob_table.rows[0].cells[1].text = 'Solução'
    
    prob_table.rows[1].cells[0].text = 'Erro ao baixar dados'
    prob_table.rows[1].cells[1].text = 'Verificar conexão com internet e validade do ticker'
    
    prob_table.rows[2].cells[0].text = 'DataFrame vazio'
    prob_table.rows[2].cells[1].text = 'Ajustar período de datas ou verificar se há dados disponíveis'
    
    prob_table.rows[3].cells[0].text = 'Import error'
    prob_table.rows[3].cells[1].text = 'Verificar se ambiente virtual está ativo e dependências instaladas'
    
    doc.add_paragraph()
    
    # ==================== RODAPÉ ====================
    doc.add_page_break()
    adicionar_titulo(doc, 'Informações Adicionais', 1)
    
    doc.add_paragraph('Projeto: QT - Análise Quantitativa de Mercado Financeiro')
    doc.add_paragraph(f'Data de Criação: Janeiro 2026')
    doc.add_paragraph(f'Última Atualização: {datetime.now().strftime("%d/%m/%Y")}')
    doc.add_paragraph('Versão: 0.1.0 (Alpha)')
    
    doc.add_paragraph()
    adicionar_paragrafo(doc, 
        'Esta documentação foi gerada automaticamente e deve ser atualizada '
        'conforme o projeto evolui. Mantenha este documento sincronizado com '
        'as mudanças no código para facilitar futuras manutenções.',
        negrito=False
    )
    
    # Salvar documento
    caminho_arquivo = r'c:\Users\angel\OneDrive\Documentos\Documentos\QT\Documentacao_Projeto_QT.docx'
    doc.save(caminho_arquivo)
    print(f"✅ Documentação gerada com sucesso!")
    print(f"📄 Arquivo salvo em: {caminho_arquivo}")
    
    return caminho_arquivo

if __name__ == "__main__":
    criar_documentacao()
